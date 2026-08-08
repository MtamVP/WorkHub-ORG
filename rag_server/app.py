import os
import re
import gc
import json
import glob
import time
import csv
import math
import logging
import requests
import urllib.request
import urllib.error
from typing import Dict, List, Optional, Any
from pathlib import Path

import numpy as np
from rank_bm25 import BM25Okapi
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from dotenv import load_dotenv

from fastapi import FastAPI, HTTPException, BackgroundTasks, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("WorkHub-RAG")

env_path = os.path.join(os.path.dirname(__file__), ".env")
if os.path.exists(env_path):
    load_dotenv(env_path)
load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL", "https://gqsbsqaxzpzcloaopzvv.supabase.co")
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "sb_publishable_sl9uOpcIzfzN9NZ5D_ZdsQ_FQZchyUR")
BRONZE_BUCKET = os.getenv("BRONZE_BUCKET", "general_bucket")
LOCAL_CORPUS_DIR = os.getenv("LOCAL_CORPUS_DIR", "./data_bronze")
PORT = int(os.getenv("PORT", "7860"))
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

try:
    from underthesea import word_tokenize
    HAVE_UNDERTHESEA = True
except ImportError:
    HAVE_UNDERTHESEA = False

try:
    from pypdf import PdfReader
    HAVE_PYPDF = True
except ImportError:
    HAVE_PYPDF = False

try:
    import docx
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

try:
    import openpyxl
    HAVE_OPENPYXL = True
except ImportError:
    HAVE_OPENPYXL = False

gemini_model = None
HAVE_GEMINI = False

def call_gemini_llm(prompt: str, api_key: Optional[str] = None) -> Optional[str]:
    key = api_key or GEMINI_API_KEY
    if not key:
        return None

    model_candidates = [
        "gemini-2.5-flash",
        "gemini-2.0-flash",
        "gemini-2.5-pro",
        "gemini-flash-latest",
        "gemini-pro-latest",
        "gemini-2.0-flash-lite"
    ]

    for m_name in model_candidates:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{m_name}:generateContent?key={key}"
            payload = {
                "contents": [
                    {
                        "parts": [{"text": prompt}]
                    }
                ],
                "generationConfig": {
                    "temperature": 0.2,
                    "topP": 0.95,
                    "maxOutputTokens": 8192
                }
            }
            resp = requests.post(url, json=payload, timeout=60)
            if resp.status_code == 200:
                data = resp.json()
                candidates = data.get("candidates", [])
                if candidates:
                    parts = candidates[0].get("content", {}).get("parts", [])
                    if parts and "text" in parts[0]:
                        ans_text = parts[0]["text"].strip()
                        if len(ans_text) > 10:
                            logger.info(f"Gemini REST success with model '{m_name}'.")
                            return ans_text
            else:
                logger.warning(f"Gemini REST '{m_name}' status {resp.status_code}: {resp.text[:200]}")
        except Exception as e:
            logger.warning(f"Gemini REST attempt '{m_name}' failed: {e}")
            continue

    try:
        import google.generativeai as genai
        genai.configure(api_key=key)
        for m_name in model_candidates:
            try:
                model = genai.GenerativeModel(m_name)
                resp = model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        max_output_tokens=8192,
                        temperature=0.2,
                        top_p=0.95
                    )
                )
                if resp and resp.text and resp.text.strip():
                    logger.info(f"Gemini SDK success with model '{m_name}'.")
                    return resp.text.strip()
            except Exception as ex:
                logger.warning(f"Gemini SDK candidate '{m_name}' failed: {ex}")
                continue
    except Exception as e:
        logger.warning(f"Error executing Gemini SDK: {e}")

    return None

if GEMINI_API_KEY:
    try:
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        HAVE_GEMINI = True
        logger.info("Gemini API configured successfully from env.")
    except Exception as e:
        logger.warning(f"Failed to configure Gemini: {e}")
        HAVE_GEMINI = False


def clean_document_title(raw_name: str) -> str:
    if not raw_name:
        return ""
    base = raw_name.split("#")[0]
    clean = re.sub(r'^(?:F_)?\d{10,20}_?', '', base, flags=re.IGNORECASE)
    root, ext = os.path.splitext(clean)
    
    subs = [
        (r'(?i)\b(?:N_I_DUNG|NOI_DUNG)\b', 'Nội Dung'),
        (r'(?i)\b(?:KH_A_H_C|KHOA_HOC)\b', 'Khóa Học'),
        (r'(?i)\b(?:C_A|CUA)\b', 'Của'),
        (r'(?i)\b(?:B_O_C_O|BAO_CAO)\b', 'Báo Cáo'),
        (r'(?i)\b(?:QUY__NH|QUY_DINH)\b', 'Quy Định'),
        (r'(?i)\b(?:T_I_LI_U|TAI_LIEU)\b', 'Tài Liệu'),
        (r'(?i)\b(?:T_I_CH_NH|TAI_CHINH)\b', 'Tài Chính'),
        (r'(?i)\b(?:H_P___NG|HOP_DONG)\b', 'Hợp Đồng'),
        (r'(?i)\b(?:K_HO_CH|KE_HOACH)\b', 'Kế Hoạch'),
        (r'(?i)\b(?:TH_NG_B_O|THONG_BAO)\b', 'Thông Báo'),
        (r'(?i)\b(?:H__NG_D_N|HUONG_DAN)\b', 'Hướng Dẫn'),
        (r'(?i)\b(?:QUY_TR_NH|QUY_TRINH)\b', 'Quy Trình'),
        (r'(?i)\b(?:CH_NH_S_CH|CHINH_SACH)\b', 'Chính Sách'),
        (r'(?i)\b(?:DOANH_NGHIEP)\b', 'Doanh Nghiệp'),
        (r'(?i)\b(?:PHAN_TICH)\b', 'Phân Tích'),
    ]
    for pattern, repl in subs:
        root = re.sub(pattern, repl, root)
    
    root = re.sub(r'_+', ' ', root).strip()
    root = re.sub(r'\s+', ' ', root)
    
    if not root:
        root = base
    return f"{root}{ext}".strip()


SYNONYM_EXPANSIONS = {
    "khoa hoc": ["khoá học", "khóa học", "đào tạo", "chương trình", "bài giảng", "syllabus", "nội dung", "module", "bài học", "học phần", "chuyên đề"],
    "lo trinh": ["lộ trình", "kế hoạch", "các giai đoạn", "quy trình", "roadmap", "timeline", "thời gian biểu", "tiến độ", "các bước", "hướng dẫn"],
    "chi phi": ["chi phí", "học phí", "giá", "ngân sách", "bảng giá", "thanh toán", "ưu đãi", "học bổng", "tiền", "đơn giá", "kinh phí"],
    "tom tat": ["tóm tắt", "tổng quan", "mục lục", "khái quát", "sơ lược", "điểm chính", "kết luận", "toàn văn", "tất cả", "nội dung chính"],
    "quy dinh": ["quy định", "chính sách", "điều khoản", "nội quy", "nghĩa vụ", "yêu cầu", "điều kiện", "tiêu chuẩn", "quy chế", "nguyên tắc"],
    "giang vien": ["giảng viên", "người hướng dẫn", "mentor", "thầy cô", "chuyên gia", "tác giả", "diễn giả", "huấn luyện viên"],
    "thuc hanh": ["thực hành", "bài tập", "dự án", "project", "lab", "bài kiểm tra", "đánh giá", "ứng dụng", "case study", "thực chiến"],
    "muc tieu": ["mục tiêu", "chuẩn đầu ra", "kết quả đạt được", "lợi ích", "kiến thức", "kỹ năng", "đầu ra", "kỳ vọng"],
    "tai chinh": ["tài chính", "doanh thu", "lợi nhuận", "chi tiêu", "dòng tiền", "báo cáo", "kế toán", "ngân sách", "hóa đơn", "thu chi"],
    "hop dong": ["hợp đồng", "thỏa thuận", "ký kết", "bên a", "bên b", "phụ lục", "cam kết", "nghĩa vụ", "thanh lý"],
    "bao cao": ["báo cáo", "thống kê", "tổng kết", "đánh giá", "kết quả", "tình hình", "biên bản", "nghiệm thu"],
    "nhan su": ["nhân sự", "nhân viên", "tuyển dụng", "nhân sự", "chức vụ", "phòng ban", "đãi ngộ", "lương thưởng", "phúc lợi"],
    "ky nang": ["kỹ năng", "năng lực", "kinh nghiệm", "chuyên môn", "kiến thức", "thành thạo", "hiểu biết"],
    "cong cu": ["công cụ", "phần mềm", "hệ thống", "nền tảng", "framework", "ngôn ngữ", "thư viện", "công nghệ"]
}

def remove_vietnamese_accents(text: str) -> str:
    s = text.lower()
    s = re.sub(r'[àáạảãâầấậẩẫăằắặẳẵ]', 'a', s)
    s = re.sub(r'[èéẹẻẽêềếệểễ]', 'e', s)
    s = re.sub(r'[ìíịỉĩ]', 'i', s)
    s = re.sub(r'[òóọỏõôồốộổỗơờớợởỡ]', 'o', s)
    s = re.sub(r'[ùúụủũưừứựửữ]', 'u', s)
    s = re.sub(r'[ỳýỵỷỹ]', 'y', s)
    s = re.sub(r'đ', 'd', s)
    return s

def expand_vietnamese_query(query: str) -> str:
    q_norm = remove_vietnamese_accents(query)
    added_terms = []
    for key, syns in SYNONYM_EXPANSIONS.items():
        if key in q_norm:
            added_terms.extend(syns[:4])
    if added_terms:
        unique_syns = list(dict.fromkeys(added_terms))
        return f"{query} {' '.join(unique_syns)}"
    return query

def tokenize_vietnamese(text: str) -> List[str]:
    clean_text = re.sub(r'[^\w\s]', " ", text).lower()
    if HAVE_UNDERTHESEA:
        try:
            tokens = word_tokenize(clean_text, format="text").split()
            return tokens
        except Exception:
            pass
    return clean_text.split()


def extract_text_from_file(file_path: str) -> List[Dict[str, Any]]:
    ext = os.path.splitext(file_path)[1].lower()
    doc_id = os.path.basename(file_path)
    chunks = []

    if ext == ".pdf" and HAVE_PYPDF:
        try:
            reader = PdfReader(file_path)
            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    chunks.append({
                        "doc_id": doc_id,
                        "chunk_id": f"{doc_id}#Trang_{page_idx + 1}",
                        "page": page_idx + 1,
                        "text": text
                    })
        except Exception as e:
            logger.warning(f"Error reading PDF {file_path}: {e}")

    elif ext in [".docx", ".doc"] and HAVE_DOCX:
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
            
            combined = "\n\n".join(full_text)
            if combined:
                chunks.append({
                    "doc_id": doc_id,
                    "chunk_id": f"{doc_id}#Toàn_văn",
                    "text": combined
                })
        except Exception as e:
            logger.warning(f"Error reading DOCX {file_path}: {e}")

    elif ext in [".xlsx", ".xls"] and HAVE_OPENPYXL:
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows_text = []
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                    if row_vals:
                        rows_text.append(" , ".join(row_vals))
                if rows_text:
                    chunks.append({
                        "doc_id": doc_id,
                        "chunk_id": f"{doc_id}#Sheet_{sheet}",
                        "text": f"Bảng tính {sheet}:\n" + "\n".join(rows_text)
                    })
        except Exception as e:
            logger.warning(f"Error reading Excel {file_path}: {e}")

    elif ext == ".csv":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                rows = [" | ".join(r) for r in reader if any(r)]
                if rows:
                    chunks.append({
                        "doc_id": doc_id,
                        "chunk_id": f"{doc_id}#CSV",
                        "text": "\n".join(rows)
                    })
        except Exception as e:
            logger.warning(f"Error reading CSV {file_path}: {e}")

    elif ext == ".json":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw = json.load(f)
                if isinstance(raw, list):
                    for idx, item in enumerate(raw):
                        if isinstance(item, dict):
                            t = item.get("passage") or item.get("content") or item.get("text") or json.dumps(item, ensure_ascii=False)
                            c_id = str(item.get("id", f"{doc_id}_{idx}"))
                            chunks.append({"doc_id": doc_id, "chunk_id": c_id, "text": str(t)})
                elif isinstance(raw, dict):
                    for k, v in raw.items():
                        if isinstance(v, str) and v.strip():
                            chunks.append({"doc_id": doc_id, "chunk_id": str(k), "text": v})
                        elif isinstance(v, dict):
                            t = v.get("passage") or v.get("content") or v.get("text") or json.dumps(v, ensure_ascii=False)
                            chunks.append({"doc_id": doc_id, "chunk_id": str(k), "text": str(t)})
        except Exception as e:
            logger.warning(f"Error reading JSON {file_path}: {e}")

    else:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read().strip()
                if text:
                    chunks.append({
                        "doc_id": doc_id,
                        "chunk_id": f"{doc_id}#Raw",
                        "text": text
                    })
        except Exception as e:
            logger.warning(f"Error reading text {file_path}: {e}")

    return chunks


def adaptive_chunking(doc_id: str, text: str, chunk_size: int = 700, chunk_overlap: int = 150) -> List[Dict[str, str]]:
    legal_segments = re.split(r'(?i)\n(Điều\s+\d+.*?)(?=\nĐiều\s+\d+|$)', text)
    if len(legal_segments) > 2:
        chunks = []
        for seg in legal_segments:
            seg = seg.strip()
            if len(seg) > 30:
                chunks.append({
                    "doc_id": doc_id,
                    "doc_text": seg
                })
        if chunks:
            return chunks

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    curr_chunk = []
    curr_len = 0

    for p in paragraphs:
        p_len = len(p)
        if curr_len + p_len > chunk_size and curr_chunk:
            combined = "\n\n".join(curr_chunk)
            chunks.append({
                "doc_id": doc_id,
                "doc_text": combined
            })
            curr_chunk = [p]
            curr_len = p_len
        else:
            curr_chunk.append(p)
            curr_len += p_len

    if curr_chunk:
        chunks.append({
            "doc_id": doc_id,
            "doc_text": "\n\n".join(curr_chunk)
        })

    return chunks if chunks else [{"doc_id": doc_id, "doc_text": text}]


def split_sentences(text: str) -> List[str]:
    raw_sentences = re.split(r'(?<=[.!?\n])\s+', text)
    sentences = [s.strip() for s in raw_sentences if len(s.strip()) > 15]
    return sentences


def build_global_idf(corpus_tokens: List[List[str]]) -> Dict[str, float]:
    num_docs = len(corpus_tokens)
    if num_docs == 0:
        return {}
    df = {}
    for doc in corpus_tokens:
        seen = set(doc)
        for token in seen:
            df[token] = df.get(token, 0) + 1

    idf = {}
    for token, freq in df.items():
        idf[token] = math.log((num_docs - freq + 0.5) / (freq + 0.5) + 1.0)
    return idf


def extract_best_sentences(query: str, document_text: str, global_idf: Dict[str, float], top_k: int = 25, max_chars: int = 6000) -> str:
    sentences = split_sentences(document_text)
    if not sentences:
        return document_text[:max_chars]

    q_tokens = tokenize_vietnamese(query)
    q_token_set = set(q_tokens)
    if not q_token_set:
        return document_text[:max_chars]

    sent_scores = []
    for idx, sent in enumerate(sentences):
        s_tokens = tokenize_vietnamese(sent)
        score = sum(global_idf.get(t, 1.0) for t in s_tokens if t in q_token_set)
        sent_scores.append((score, idx, sent))

    sent_scores.sort(key=lambda x: x[0], reverse=True)
    selected_indices = sorted([item[1] for item in sent_scores[:top_k] if item[0] > 0])

    if not selected_indices:
        return "\n".join(sentences[:top_k])[:max_chars]

    result = []
    cur_len = 0
    for idx in selected_indices:
        s = sentences[idx]
        if cur_len + len(s) > max_chars:
            break
        result.append(s)
        cur_len += len(s)

    return "\n\n".join(result) if result else document_text[:max_chars]


class BronzeStorageManager:
    def __init__(self, local_dir: str = LOCAL_CORPUS_DIR):
        self.local_dir = local_dir
        os.makedirs(self.local_dir, exist_ok=True)
        self.headers = {
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {SUPABASE_KEY}",
            "Content-Type": "application/json"
        }

    def sync_from_supabase(self, bucket_name: str = BRONZE_BUCKET) -> List[str]:
        downloaded = []
        active_remote_files = set()

        if not SUPABASE_URL or not SUPABASE_KEY:
            logger.warning("Supabase URL or Key missing.")
            return downloaded

        target_buckets = ["general_bucket", "finance_bucket", "science_bucket", "bronze_storage"]
        if bucket_name and bucket_name not in target_buckets:
            target_buckets.append(bucket_name)

        try:
            tbl_url = f"{SUPABASE_URL}/rest/v1/files?select=*&deleted_at=is.null"
            t_req = urllib.request.Request(tbl_url, headers=self.headers)
            with urllib.request.urlopen(t_req, timeout=15) as t_resp:
                records = json.loads(t_resp.read().decode("utf-8"))
            for rec in records:
                name = rec.get("name")
                storage_path = rec.get("storage_path") or ""
                if not storage_path or not name:
                    continue
                parts = storage_path.split("/", 1)
                b_name = parts[0]
                inner = parts[1] if len(parts) > 1 else name
                
                safe_name = os.path.basename(storage_path) or name.replace("/", "_").replace("\\", "_")
                active_remote_files.add(safe_name)

                local_path = os.path.join(self.local_dir, safe_name)
                if not os.path.exists(local_path):
                    try:
                        d_req = urllib.request.Request(f"{SUPABASE_URL}/storage/v1/object/public/{b_name}/{inner}", headers=self.headers)
                        with urllib.request.urlopen(d_req, timeout=20) as d_resp:
                            data = d_resp.read()
                            with open(local_path, "wb") as out_f:
                                out_f.write(data)
                            if safe_name not in downloaded:
                                downloaded.append(safe_name)
                                logger.info(f"Downloaded from files table: {safe_name}")
                    except Exception as ex:
                        logger.warning(f"Error downloading {storage_path}: {ex}")
                else:
                    if safe_name not in downloaded:
                        downloaded.append(safe_name)
        except Exception as e:
            logger.warning(f"Error querying files table: {e}")

        for b_name in target_buckets:
            for prefix in ["bronze", ""]:
                try:
                    list_url = f"{SUPABASE_URL}/storage/v1/object/list/{b_name}"
                    req = urllib.request.Request(
                        list_url,
                        data=json.dumps({"prefix": prefix, "limit": 100, "offset": 0}).encode("utf-8"),
                        headers=self.headers,
                        method="POST"
                    )
                    with urllib.request.urlopen(req, timeout=15) as resp:
                        file_list = json.loads(resp.read().decode("utf-8"))

                    for f in file_list:
                        fname = f.get("name")
                        if not fname or fname.startswith(".") or fname == "bronze":
                            continue
                        
                        active_remote_files.add(fname)
                        local_path = os.path.join(self.local_dir, fname)

                        if not os.path.exists(local_path):
                            full_key = f"{prefix}/{fname}".strip("/") if prefix else fname
                            pub_url = f"{SUPABASE_URL}/storage/v1/object/public/{b_name}/{full_key}"
                            auth_url = f"{SUPABASE_URL}/storage/v1/object/authenticated/{b_name}/{full_key}"
                            
                            file_bytes = None
                            try:
                                d_req = urllib.request.Request(pub_url, headers=self.headers)
                                with urllib.request.urlopen(d_req, timeout=20) as d_resp:
                                    file_bytes = d_resp.read()
                            except Exception:
                                try:
                                    d_req = urllib.request.Request(auth_url, headers=self.headers)
                                    with urllib.request.urlopen(d_req, timeout=20) as d_resp:
                                        file_bytes = d_resp.read()
                                except Exception:
                                    pass
                            
                            if file_bytes:
                                with open(local_path, "wb") as out_f:
                                    out_f.write(file_bytes)
                                if fname not in downloaded:
                                    downloaded.append(fname)
                                    logger.info(f"Downloaded: {fname} ({len(file_bytes)} bytes)")
                        else:
                            if fname not in downloaded:
                                downloaded.append(fname)
                except Exception as ex:
                    logger.warning(f"Error scanning {b_name}/{prefix}: {ex}")

        existing_local_files = glob.glob(os.path.join(self.local_dir, "*.*"))
        for loc_f in existing_local_files:
            bname = os.path.basename(loc_f)
            if bname not in active_remote_files:
                try:
                    os.remove(loc_f)
                    logger.info(f"Purged deleted file from disk: {bname}")
                except Exception as ex:
                    logger.warning(f"Failed to remove orphaned file {loc_f}: {ex}")

        return list(active_remote_files)

    def get_document_catalog(self) -> List[Dict[str, Any]]:
        catalog = []
        all_files = glob.glob(os.path.join(self.local_dir, "*.*"))
        seen_files = set()

        for file_path in all_files:
            fname = os.path.basename(file_path)
            if fname in seen_files or fname.startswith("."):
                continue
            seen_files.add(fname)

            size_bytes = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            if size_bytes >= 1024 * 1024:
                size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
            elif size_bytes >= 1024:
                size_str = f"{size_bytes / 1024:.1f} KB"
            else:
                size_str = f"{size_bytes} B"

            chunks = extract_text_from_file(file_path)
            total_pages = len(chunks)

            preview_text = ""
            if chunks:
                preview_text = "\n".join([c["text"] for c in chunks[:2]]).strip()
                if len(preview_text) > 300:
                    preview_text = preview_text[:300] + "..."

            display_title = clean_document_title(fname)

            catalog.append({
                "file_name": fname,
                "display_name": display_title,
                "clean_name": re.sub(r'\.[^/.]+$', '', display_title),
                "total_pages": total_pages,
                "file_size": size_str,
                "preview": preview_text or "Tài liệu văn bản số hóa"
            })

        return catalog

    def load_all_documents(self) -> Dict[str, str]:
        corpus: Dict[str, str] = {}
        all_files = glob.glob(os.path.join(self.local_dir, "*.*"))
        
        for file_path in all_files:
            extracted_chunks = extract_text_from_file(file_path)
            for c in extracted_chunks:
                corpus[c["chunk_id"]] = c["text"]

        logger.info(f"Loaded {len(corpus)} document chunks from {len(all_files)} files.")
        return corpus


class RAGEngine:
    def __init__(self):
        self.corpus: Dict[str, str] = {}
        self.doc_ids: List[str] = []
        self.article_chunks: List[Dict[str, str]] = []
        self.file_to_chunks: Dict[str, List[str]] = {}
        self.bm25_index: Optional[BM25Okapi] = None
        self.tfidf_vectorizer: Optional[TfidfVectorizer] = None
        self.tfidf_matrix = None
        self.global_idf: Dict[str, float] = {}
        self.is_ready = False
        self.last_indexed_time = None

    def build_index(self, corpus: Dict[str, str]):
        if not corpus:
            self.corpus = {}
            self.doc_ids = []
            self.article_chunks = []
            self.file_to_chunks = {}
            self.bm25_index = None
            self.tfidf_vectorizer = None
            self.tfidf_matrix = None
            self.global_idf = {}
            self.is_ready = False
            self.last_indexed_time = None
            logger.info("RAG Index cleared (0 documents).")
            return

        self.corpus = corpus
        self.doc_ids = list(corpus.keys())

        self.file_to_chunks = {}
        for c_id in self.doc_ids:
            parent_file = c_id.split("#")[0]
            if parent_file not in self.file_to_chunks:
                self.file_to_chunks[parent_file] = []
            self.file_to_chunks[parent_file].append(c_id)

        self.article_chunks = []
        for doc_id, doc_text in self.corpus.items():
            self.article_chunks.extend(adaptive_chunking(doc_id, doc_text))

        corpus_tokens = [tokenize_vietnamese(text) for text in self.corpus.values()]
        self.bm25_index = BM25Okapi(corpus_tokens)
        self.global_idf = build_global_idf(corpus_tokens)

        chunk_texts = [c["doc_text"] for c in self.article_chunks]
        self.tfidf_vectorizer = TfidfVectorizer(ngram_range=(1, 2), max_features=30000)
        self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(chunk_texts)

        self.is_ready = True
        self.last_indexed_time = time.strftime("%Y-%m-%d %H:%M:%S")
        logger.info(f"RAG Index ready: {len(self.corpus)} chunks indexed across {len(self.file_to_chunks)} files.")

    def query(
        self,
        question: str,
        top_k_docs: int = 5,
        use_llm: bool = True,
        gemini_api_key: Optional[str] = None,
        history: Optional[List[Dict[str, str]]] = None,
        catalog: Optional[List[Dict[str, Any]]] = None,
        focus_doc_id: Optional[str] = None
    ) -> Dict[str, Any]:
        if not self.is_ready or not self.corpus:
            history_text = ""
            if history and isinstance(history, list) and len(history) > 0:
                hist_lines = []
                for h in history[-6:]:
                    r = "Người dùng" if h.get("role") == "user" else "Ciel (AI)"
                    hist_lines.append(f"{r}: {h.get('content', '')[:350]}")
                history_text = "LỊCH SỬ TRÒ CHUYỆN GẦN ĐÂY:\n" + "\n".join(hist_lines)

            general_prompt = f"""Bạn là Ciel - Siêu Trí tuệ Nhân tạo & Cố vấn Tri thức cao cấp của WorkHub.
Hiện tại kho lưu trữ Bronze Storage đang trống (chưa có tài liệu nào được tải lên).

{history_text}

[CÂU HỎI CỦA NGƯỜI DÙNG]:
{question}

[QUY TẮC PHẢN HỒI]:
1. **DÒNG ĐẦU TIÊN**: Luôn bắt đầu bằng dòng ghi chú sau:
> 💡 *Lưu ý: Do kho Bronze Storage hiện chưa có tài liệu nào, Ciel sẽ giải đáp dựa trên tri thức tổng quát.*

2. **TRẢ LỜI ĐẲNG CẤP & CHI TIẾT**:
   - Trả lời thẳng vào trọng tâm câu hỏi, phân tích chuyên sâu, mạch lạc và đầy đủ.
   - Trình bày định dạng Markdown đẹp mắt (tiêu đề ###, gạch đầu dòng, in đậm thuật ngữ, bảng biểu nếu có).
   
3. **GỢI Ý CÂU HỎI TIẾP THEO**: Ở dòng cuối cùng, tạo đúng định dạng:
[GỢI Ý CÂU HỎI TIẾP THEO]:
- Câu hỏi gợi ý 1?
- Câu hỏi gợi ý 2?
- Câu hỏi gợi ý 3?"""

            ans = call_gemini_llm(general_prompt, api_key=gemini_api_key)
            if not ans:
                ans = "> 💡 *Lưu ý: Do kho Bronze Storage hiện chưa có tài liệu nào, Ciel sẽ giải đáp dựa trên tri thức tổng quát.*\n\nXin chào! Tôi là Ciel. Hiện tại kho tài liệu chưa có dữ liệu nào được lập chỉ mục. Bạn có thể tải tài liệu lên Bronze Storage để tôi phân tích, hoặc đặt bất kỳ câu hỏi nào để tôi hỗ trợ nhé!"

            suggestions = []
            if "[GỢI Ý CÂU HỎI TIẾP THEO]:" in ans:
                parts = ans.split("[GỢI Ý CÂU HỎI TIẾP THEO]:")
                main_ans = parts[0].strip()
                sugg_text = parts[1].strip()
                for line in sugg_text.split("\n"):
                    clean_s = re.sub(r'^\s*[-*•\d.]+\s*', '', line).strip()
                    if clean_s and len(clean_s) > 5 and len(clean_s) < 150:
                        suggestions.append(clean_s)
                ans = main_ans

            if not suggestions:
                suggestions = [
                    "Hướng dẫn cách tải tài liệu lên Bronze Storage",
                    "WorkHub có những tính năng nào hỗ trợ công việc?",
                    "Lập kế hoạch làm việc và quản lý thời gian hiệu quả"
                ]

            return {
                "answer": ans,
                "sources": [],
                "best_doc_id": "",
                "retrieved_docs": [],
                "suggestions": suggestions[:3]
            }

        expanded_query = expand_vietnamese_query(question)
        q_tokens = tokenize_vietnamese(expanded_query)
        q_norm = remove_vietnamese_accents(question)

        is_global_inquiry = any(kw in q_norm for kw in [
            "tom tat", "tong quan", "toan bo", "tat ca", "co nhung tai lieu", "danh sach file",
            "kho co gi", "gioi thieu", "bao gom nhung gi", "so sanh"
        ])

        bm25_scores = self.bm25_index.get_scores(q_tokens)
        bm25_ranks = {
            self.doc_ids[i]: rank + 1
            for rank, i in enumerate(np.argsort(bm25_scores)[::-1][:30])
        }

        q_vec = self.tfidf_vectorizer.transform([expanded_query])
        cos_sim = cosine_similarity(q_vec, self.tfidf_matrix)[0]
        top_chunk_indices = np.argsort(cos_sim)[::-1][:30]

        tfidf_ranks = {}
        for rank, chunk_idx in enumerate(top_chunk_indices):
            parent_doc_id = self.article_chunks[chunk_idx]["doc_id"]
            if parent_doc_id not in tfidf_ranks:
                tfidf_ranks[parent_doc_id] = rank + 1

        title_boost = {}
        for doc_id in self.doc_ids:
            parent_file = doc_id.split("#")[0]
            clean_name = clean_document_title(parent_file)
            clean_norm = remove_vietnamese_accents(clean_name)
            
            words = [w for w in clean_norm.split() if len(w) > 2]
            match_count = sum(1 for w in words if w in q_norm)
            if match_count >= 2:
                title_boost[doc_id] = 0.05 * match_count
            elif match_count == 1:
                title_boost[doc_id] = 0.02

        fused_score = {}
        for doc_id in set(bm25_ranks) | set(tfidf_ranks) | set(title_boost):
            bm25_r = bm25_ranks.get(doc_id, 999)
            tfidf_r = tfidf_ranks.get(doc_id, 999)
            score = (
                1.0 / (60.0 + bm25_r) +
                1.0 / (60.0 + tfidf_r) +
                title_boost.get(doc_id, 0.0)
            )
            doc_raw = self.corpus.get(doc_id, "")
            doc_norm = remove_vietnamese_accents(doc_raw)
            for key_term in [w for w in q_norm.split() if len(w) > 2]:
                if key_term in doc_norm:
                    score += 0.003
            fused_score[doc_id] = score

        ranked_doc_ids = [
            doc_id
            for doc_id, _ in sorted(fused_score.items(), key=lambda x: (x[1], x[0]), reverse=True)[:top_k_docs]
        ]

        if focus_doc_id and focus_doc_id.strip() and focus_doc_id.strip().lower() != "all":
            clean_focus = focus_doc_id.strip()
            focus_chunks = self.file_to_chunks.get(clean_focus, [])
            if not focus_chunks:
                for f_key in self.file_to_chunks.keys():
                    if clean_focus.lower() in f_key.lower() or f_key.lower() in clean_focus.lower():
                        focus_chunks = self.file_to_chunks[f_key]
                        clean_focus = f_key
                        break

            if focus_chunks:
                ranked_doc_ids = focus_chunks[:top_k_docs]
                all_ordered_files = [clean_focus]
            else:
                active_parent_files = list(dict.fromkeys([d.split("#")[0] for d in ranked_doc_ids]))
                other_files = [f for f in self.file_to_chunks.keys() if f not in active_parent_files]
                all_ordered_files = active_parent_files + other_files
        else:
            active_parent_files = list(dict.fromkeys([d.split("#")[0] for d in ranked_doc_ids]))
            other_files = [f for f in self.file_to_chunks.keys() if f not in active_parent_files]
            all_ordered_files = active_parent_files + other_files

        context_blocks = []
        for p_file in all_ordered_files:
            file_chunks = self.file_to_chunks.get(p_file, [])
            clean_file = clean_document_title(p_file)
            
            file_full_text = []
            for c_id in file_chunks:
                c_text = self.corpus.get(c_id, "").strip()
                if c_text:
                    parts = c_id.split("#")
                    page_info = f"Trang {parts[1].replace('_', ' ').replace('Trang ', '')}" if len(parts) > 1 else "Nội dung"
                    file_full_text.append(f"--- [Phần: {page_info}] ---\n{c_text}")
            
            if file_full_text:
                context_blocks.append(f"==================================================\n📁 [TÀI LIỆU TOÀN VĂN: {clean_file}]\n==================================================\n" + "\n\n".join(file_full_text))

        combined_context = "\n\n".join(context_blocks)

        extracted_sections = []
        for d_id in ranked_doc_ids:
            raw_text = self.corpus.get(d_id, "")
            ext = extract_best_sentences(question, raw_text, self.global_idf, top_k=25, max_chars=4500)
            if ext:
                parts = d_id.split("#")
                clean_file = clean_document_title(parts[0])
                page_info = f" ({parts[1].replace('_', ' ')})" if len(parts) > 1 else ""
                extracted_sections.append(f"#### 📄 Trích từ: {clean_file}{page_info}\n{ext}")

        fallback_answer = "\n\n---\n\n".join(extracted_sections) if extracted_sections else "Không tìm thấy nội dung chi tiết trong tài liệu."
        final_answer = fallback_answer

        catalog_summary = ""
        if catalog:
            cat_lines = [f"- 📁 **{c.get('display_name')}** ({c.get('total_pages', 1)} trang, {c.get('file_size', '')})" for c in catalog]
            catalog_summary = "DANH MỤC CÁC TÀI LIỆU HIỆN CÓ TRONG KHO:\n" + "\n".join(cat_lines)

        history_text = ""
        if history and isinstance(history, list) and len(history) > 0:
            hist_lines = []
            for h in history[-6:]:
                r = "Người dùng" if h.get("role") == "user" else "Ciel (AI)"
                hist_lines.append(f"{r}: {h.get('content', '')[:350]}")
            history_text = "LỊCH SỬ TRÒ CHUYỆN GẦN ĐÂY:\n" + "\n".join(hist_lines)

        prompt = f"""Bạn là Ciel - Siêu Trí tuệ Nhân tạo & Cố vấn Phân tích Tri thức cao cấp của WorkHub.
Dưới đây là TOÀN BỘ dữ liệu tài liệu nội bộ đã được số hóa từ kho lưu trữ Bronze Storage.

{catalog_summary}

{history_text}

[KHO TRI THỨC TOÀN VĂN]:
{combined_context[:120000]}

[CÂU HỎI CỦA NGƯỜI DÙNG]:
{question}

[QUY TẮC TRẢ LỜI ĐẲNG CẤP & CHUYÊN NGHIỆP]:
1. **TRẢ LỜI TRỰC DIỆN & CHÍNH XÁC NHẤT**: Hãy trả lời thẳng vào trọng tâm câu hỏi của người dùng ngay từ phần đầu tiên, giải thích rõ ràng, mạch lạc và súc tích.
2. **CHI TIẾT, ĐẦY ĐỦ & TOÀN DIỆN**:
   - Trích xuất toàn bộ thông tin chi tiết, lộ trình, các module, điều khoản, tiêu chuẩn, số liệu và các bước thực hiện có trong tài liệu.
   - Tuyệt đối KHÔNG trả lời ngắn ngủi, chung chung hoặc cắt bớt chi tiết. Nếu tài liệu có 5 module, hãy phân tích đầy đủ cả 5 module.
3. **TRÌNH BÀY ĐẸP MẮT (MARKDOWN)**:
   - Sử dụng tiêu đề (###, ####) phân cấp nội dung logic.
   - In đậm (**từ khóa, thuật ngữ**) quan trọng.
   - Sử dụng bảng biểu Markdown (`| Cột 1 | Cột 2 |`) nếu có danh sách khóa học, bảng giá, so sánh hoặc mốc tiến độ.
   - Gạch đầu dòng rõ ràng, dễ theo dõi.
4. **TRÍCH DẪN & XỬ LÝ CÂU HỎI NGOÀI TÀI LIỆU**:
   - Nếu câu hỏi liên quan đến tài liệu: Nêu rõ thông tin được trích xuất từ tài liệu nào và số trang tương ứng.
   - Nếu câu hỏi là kiến thức chung, chào hỏi, toán, lập trình hoặc không có trong tài liệu: Hãy trả lời đầy đủ, chi tiết và thông minh dựa trên tri thức tổng quát của bạn (kèm ghi chú: `> 💡 *Lưu ý: Nội dung này được giải đáp từ tri thức tổng quát do không có trong tài liệu lưu trữ.*`).
5. **GỢI Ý CÂU HỎI TIẾP THEO**: Ở dòng cuối cùng của câu trả lời, hãy tạo đúng định dạng sau với 3 câu hỏi gợi ý thông minh nhất để người dùng khám phá sâu hơn:
[GỢI Ý CÂU HỎI TIẾP THEO]:
- Câu hỏi gợi ý 1?
- Câu hỏi gợi ý 2?
- Câu hỏi gợi ý 3?"""

        llm_res = call_gemini_llm(prompt, api_key=gemini_api_key)
        if llm_res and len(llm_res.strip()) > 30:
            final_answer = llm_res

        suggestions = []
        if "[GỢI Ý CÂU HỎI TIẾP THEO]:" in final_answer:
            parts = final_answer.split("[GỢI Ý CÂU HỎI TIẾP THEO]:")
            main_ans = parts[0].strip()
            sugg_text = parts[1].strip()
            
            for line in sugg_text.split("\n"):
                clean_s = re.sub(r'^\s*[-*•\d.]+\s*', '', line).strip()
                if clean_s and len(clean_s) > 5 and len(clean_s) < 150:
                    suggestions.append(clean_s)
            
            final_answer = main_ans
        
        if not suggestions:
            suggestions = [
                "Tóm tắt chi tiết toàn bộ nội dung của tài liệu này",
                "Có những lưu ý hoặc điều kiện quan trọng nào cần chú ý?",
                "Liệt kê các bước thực hiện hoặc lộ trình chi tiết"
            ]

        clean_sources = []
        for d_id in ranked_doc_ids:
            parts = d_id.split("#")
            clean_file = clean_document_title(parts[0])
            if len(parts) > 1:
                clean_page = parts[1].replace("_", " ")
                clean_sources.append(f"{clean_file} ({clean_page})")
            else:
                clean_sources.append(clean_file)

        retrieved_details = []
        for d_id in ranked_doc_ids:
            parts = d_id.split("#")
            clean_file = clean_document_title(parts[0])
            page_info = f" ({parts[1].replace('_', ' ')})" if len(parts) > 1 else ""
            retrieved_details.append({
                "doc_id": d_id,
                "display_name": f"{clean_file}{page_info}",
                "text": self.corpus.get(d_id, "")[:400] + "...",
                "score": float(fused_score.get(d_id, 0.0))
            })

        return {
            "answer": final_answer,
            "sources": list(dict.fromkeys(clean_sources)),
            "best_doc_id": clean_sources[0] if clean_sources else "",
            "retrieved_docs": retrieved_details,
            "suggestions": suggestions[:3]
        }


app = FastAPI(title="WorkHub Universal RAG API", version="3.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

storage_mgr = BronzeStorageManager(LOCAL_CORPUS_DIR)
rag_engine = RAGEngine()


@app.on_event("startup")
def startup_event():
    try:
        storage_mgr.sync_from_supabase()
    except Exception as e:
        logger.warning(f"Startup sync error: {e}")
    corpus = storage_mgr.load_all_documents()
    if corpus:
        rag_engine.build_index(corpus)


class ChatQueryRequest(BaseModel):
    question: str
    top_docs: Optional[int] = 5
    use_llm: Optional[bool] = True
    gemini_api_key: Optional[str] = None
    history: Optional[List[Dict[str, str]]] = []
    focus_doc_id: Optional[str] = None


class SyncRequest(BaseModel):
    bucket_name: Optional[str] = BRONZE_BUCKET


@app.get("/")
def root_endpoint():
    return {
        "service": "WorkHub Universal RAG Engine",
        "status": "online",
        "version": "3.0.0",
        "features": ["Hybrid BM25+TF-IDF", "Whole-Corpus Ingestion", "Gemini 2.5 Flash Super Reasoning", "Multi-turn Memory"]
    }


@app.get("/api/status")
def get_status():
    catalog = storage_mgr.get_document_catalog()
    return {
        "status": "ready" if rag_engine.is_ready else "empty",
        "model": "BM25 + TF-IDF + Gemini 2.5 Flash Super Reasoning",
        "device": "cpu",
        "documents_indexed": len(catalog),
        "chunks_indexed": len(rag_engine.corpus),
        "have_gemini": bool(GEMINI_API_KEY),
        "have_underthesea": HAVE_UNDERTHESEA,
        "have_pypdf": HAVE_PYPDF,
        "have_docx": HAVE_DOCX,
        "last_indexed": rag_engine.last_indexed_time,
        "local_storage_dir": LOCAL_CORPUS_DIR
    }


@app.post("/api/chat")
def chat_endpoint(req: ChatQueryRequest):
    if not req.question or not req.question.strip():
        raise HTTPException(status_code=400, detail="Câu hỏi không được để trống.")

    catalog = storage_mgr.get_document_catalog()
    should_use_llm = True if req.use_llm is None or req.use_llm is True else False

    res = rag_engine.query(
        question=req.question.strip(),
        top_k_docs=req.top_docs or 5,
        use_llm=should_use_llm,
        gemini_api_key=req.gemini_api_key,
        history=req.history or [],
        catalog=catalog,
        focus_doc_id=req.focus_doc_id
    )
    return {
        "success": True,
        "data": res
    }


@app.post("/api/sync")
def sync_bronze_storage(req: SyncRequest = SyncRequest()):
    try:
        bucket = req.bucket_name or BRONZE_BUCKET
        active_files = storage_mgr.sync_from_supabase(bucket)
        corpus = storage_mgr.load_all_documents()
        catalog = storage_mgr.get_document_catalog()
        rag_engine.build_index(corpus)
        if corpus:
            return {
                "success": True,
                "message": f"Đã đồng bộ {len(catalog)} tài liệu ({len(corpus)} trang/phân đoạn).",
                "downloaded_files": active_files,
                "total_documents": len(catalog),
                "total_chunks": len(rag_engine.article_chunks)
            }
        else:
            return {
                "success": True,
                "message": "Đã dọn dẹp sạch sẽ: Hiện không còn tài liệu nào trong kho lưu trữ.",
                "downloaded_files": [],
                "total_documents": 0,
                "total_chunks": 0
            }
    except Exception as e:
        logger.error(f"Sync error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/documents")
def list_documents():
    catalog = storage_mgr.get_document_catalog()
    return {
        "total": len(catalog),
        "items": catalog
    }


@app.get("/api/documents/page")
def get_document_page(doc_id: str):
    if not doc_id:
        raise HTTPException(status_code=400, detail="Thiếu doc_id")
    
    clean_target = doc_id.strip()
    
    if clean_target in rag_engine.corpus:
        parts = clean_target.split("#")
        clean_file = clean_document_title(parts[0])
        page_info = parts[1].replace("_", " ") if len(parts) > 1 else "Toàn văn"
        return {
            "success": True,
            "doc_id": clean_target,
            "display_name": f"{clean_file} ({page_info})",
            "text": rag_engine.corpus[clean_target]
        }
    
    for c_id, text in rag_engine.corpus.items():
        if clean_target.lower() in c_id.lower() or c_id.lower() in clean_target.lower():
            parts = c_id.split("#")
            clean_file = clean_document_title(parts[0])
            page_info = parts[1].replace("_", " ") if len(parts) > 1 else "Toàn văn"
            return {
                "success": True,
                "doc_id": c_id,
                "display_name": f"{clean_file} ({page_info})",
                "text": text
            }

    for fname, chunks in rag_engine.file_to_chunks.items():
        if clean_target.lower() in fname.lower() or fname.lower() in clean_target.lower():
            full_texts = [rag_engine.corpus.get(c, "") for c in chunks if rag_engine.corpus.get(c, "")]
            clean_file = clean_document_title(fname)
            return {
                "success": True,
                "doc_id": fname,
                "display_name": clean_file,
                "text": "\n\n--- Phân đoạn tiếp theo ---\n\n".join(full_texts)
            }

    raise HTTPException(status_code=404, detail="Không tìm thấy nội dung trang tài liệu yêu cầu.")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=PORT, reload=True)
